import re
from pathlib import Path
from typing import List, Dict, Any, Optional

import pdfplumber
from docx import Document
import pandas as pd


# -------------------------
# Normalización de líneas
# -------------------------

BULLETS = ("•", "·", "–", "-", "—", "◦", "▪", "✓", "●")

def clean_line(line: str) -> str:
    line = line.strip()
    if not line:
        return ""
    # quitar bullets/numeración tipo "1)" "1." "(1)"
    line = re.sub(r"^\(?\s*\d{1,3}\s*[\)\.\-:]\s*", "", line)
    line = line.lstrip("".join(BULLETS)).strip()
    # colapsar espacios
    line = re.sub(r"\s+", " ", line)
    return line

def split_to_candidate_lines(raw_text: str) -> List[str]:
    lines = []
    for ln in raw_text.splitlines():
        ln = clean_line(ln)
        if ln:
            lines.append(ln)
    return lines


# -------------------------
# Parse de item (cantidad/unidad/detalle)
# -------------------------

import re
from typing import Optional, Dict, Any

# Mapea unidades comunes
UNIT_MAP = {
    "unidad": "unid", "unidades": "unid", "unid": "unid", "uds": "unid", "ud": "unid",
    "pack": "pack", "paquete": "pack", "paq": "pack",
    "caja": "caja", "cajas": "caja",
    "sobre": "sobre", "sobres": "sobre",
    "bolsa": "bolsa", "bolsas": "bolsa",
    "pliego": "pliego", "pliegos": "pliego",
    "resma": "resma", "resmas": "resma",
    "frasco": "frasco", "frascos": "frasco",
}

# Asignaturas típicas (puedes ampliar)
SUBJECTS = {
    "LENGUAJE", "MATEMÁTICA", "MATEMATICA", "INGLÉS", "INGLES",
    "CIENCIAS", "CIENCIAS NATURALES", "SOCIALES",
    "ARTE Y", "TECNOLOGÍA", "TECNOLOGIA", "MÚSICA", "MUSICA",
    "TERAPIA", "MATERIALES",
}

HEADER_PATTERNS = [
    r"^LISTA DE ÚTILES",                 # LISTA DE ÚTILES...
    r"^\d+º\s+BÁSICO",                   # 3º BÁSICO
    r"^ASIGNATURA\s+ARTÍCULO",           # encabezado tabla
]
import re
from typing import List, Dict, Any

SECTION_WORDS = {
    "LENGUAJE", "MATEMÁTICA", "MATEMATICA", "INGLÉS", "INGLES",
    "CIENCIAS", "SOCIALES", "NATURALES",
    "ARTE Y", "TECNOLOGÍA", "TECNOLOGIA", "MÚSICA", "MUSICA",
    "TERAPIA", "OCUPACIONAL",
    "MATERIALES", "PARA USO", "PERSONAL", "DE ASEO",
}

def is_section_only(line: str) -> str | None:
    """Devuelve el nombre de sección si la línea es un título/Sección, si no None."""
    up = line.strip().upper()
    # normaliza múltiples espacios
    up = re.sub(r"\s+", " ", up)

    # títulos cortos tipo "SOCIALES" / "DE ASEO" / "PARA USO"
    if up in SECTION_WORDS:
        return up

    return None


def parse_lines_to_items(lines: List[str]) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    current_subject = None

    for ln in lines:
        sec = is_section_only(ln)
        if sec:
            current_subject = sec
            continue  # NO es item

        it = parse_item_line(ln)  # tu función mejorada
        if not it:
            continue

        # si detectó asignatura en la misma línea, actualiza el "current"
        if it.get("asignatura"):
            current_subject = it["asignatura"]
        else:
            it["asignatura"] = current_subject

        # ajuste unidad resma (cuando aparece como texto pero no se detectó)
        if it.get("unidad") == "unid" and re.search(r"\bresma(s)?\b", it.get("detalle",""), re.IGNORECASE):
            it["unidad"] = "resma"

        items.append(it)

    # post-proceso para unir líneas sueltas tipo "blanco"
    items = merge_loose_continuations(items)

    return items
def merge_loose_continuations(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    merged = []
    i = 0
    while i < len(items):
        cur = items[i]
        # si un item no tiene cantidad ni unidad y es una sola palabra -> probablemente continuación
        if (cur.get("cantidad") is None and cur.get("unidad") is None
            and cur.get("detalle") and len(cur["detalle"].split()) == 1
            and merged):
            # lo pegamos al anterior
            merged[-1]["detalle"] = (merged[-1]["detalle"] + " " + cur["detalle"]).strip()
            merged[-1]["item_original"] = (merged[-1]["item_original"] + " " + cur["item_original"]).strip()
            i += 1
            continue

        merged.append(cur)
        i += 1
    return merged

def is_header_line(line: str) -> bool:
    up = line.strip().upper()
    for pat in HEADER_PATTERNS:
        if re.match(pat, up):
            return True
    # líneas sueltas que son solo un nombre o fragmento de autor suelen no ser items
    if len(line.strip()) <= 2:
        return True
    return False


def parse_item_line(line: str) -> Optional[Dict[str, Any]]:
    original = line.strip()
    if not original:
        return None

    if is_header_line(original):
        return None

    # Muchos PDFs traen cosas tipo "⇒ ..." (lecturas complementarias)
    # Decide si quieres excluirlas o tratarlas como "LIBRO"
    if original.strip().startswith("⇒"):
        return None  # o manejar en otra categoría

    # 1) separar asignatura si viene pegada: "LENGUAJE Carpeta... 1"
    subject = None
    parts = original.split()
    if parts:
        first = parts[0].upper()
        # caso "ARTE Y ..." (dos tokens)
        first_two = " ".join(parts[:2]).upper() if len(parts) >= 2 else ""
        if first_two in SUBJECTS:
            subject = first_two
            rest_text = original[len(parts[0]) + 1 + len(parts[1]):].strip()
        elif first in SUBJECTS:
            subject = first
            rest_text = original[len(parts[0]):].strip()
        else:
            rest_text = original

    rest_text = rest_text.strip()

    qty = None
    unit = None
    detail = rest_text

    # 2) patrón: "... <unidad> <cantidad>" al FINAL
    m = re.search(r"\b(unidades?|uds?|ud|pack|paquete|paq|cajas?|caja|sobres?|sobre|bolsas?|bolsa|pliegos?|pliego|resmas?|resma|frascos?|frasco)\s+(\d{1,3})\s*$",
                  rest_text, flags=re.IGNORECASE)
    if m:
        unit_raw = m.group(1).lower()
        qty = int(m.group(2))
        unit = UNIT_MAP.get(unit_raw, unit_raw)
        detail = rest_text[:m.start()].strip()
    else:
        # 3) patrón: "... <cantidad>" al FINAL
        m2 = re.search(r"(\d{1,3})\s*$", rest_text)
        if m2:
            qty = int(m2.group(1))
            unit = "unid"  # por defecto
            detail = rest_text[:m2.start()].strip()

    # limpiar detalle raro
    if not detail:
        detail = rest_text

    item = {
        "item_original": original,
        "asignatura": subject,
        "cantidad": qty,
        "unidad": unit,
        "detalle": detail
    }
    return item


def parse_lines_to_items(lines: List[str]) -> List[Dict[str, Any]]:
    items = []
    for ln in lines:
        it = parse_item_line(ln)
        if it:
            items.append(it)
    return items


# -------------------------
# Extractores por tipo
# -------------------------

def extract_from_pdf(path: Path) -> str:
    text_parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text_parts.append(t)
    return "\n".join(text_parts)

def extract_from_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []

    # párrafos
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)

    # tablas (cada fila como línea)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" ".join(cells))

    return "\n".join(parts)

def extract_from_excel(path: Path) -> str:
    # Lee primera hoja; si hay varias, puedes iterar luego
    df = pd.read_excel(str(path), engine="openpyxl")
    # Convierte filas a líneas; intenta usar columnas típicas si existen
    cols = [c.lower() for c in df.columns.astype(str).tolist()]

    # heurística: si hay columna cantidad + descripción
    qty_col = None
    desc_col = None
    for i, c in enumerate(cols):
        if qty_col is None and "cant" in c:
            qty_col = df.columns[i]
        if desc_col is None and ("art" in c or "prod" in c or "desc" in c or "util" in c):
            desc_col = df.columns[i]

    lines = []
    if desc_col is not None:
        for _, row in df.iterrows():
            desc = str(row.get(desc_col, "")).strip()
            if not desc or desc.lower() == "nan":
                continue
            qty = row.get(qty_col) if qty_col is not None else None
            if pd.notna(qty):
                try:
                    qty_int = int(float(qty))
                    lines.append(f"{qty_int} {desc}")
                except Exception:
                    lines.append(desc)
            else:
                lines.append(desc)
    else:
        # fallback: unir celdas de la fila
        for _, row in df.iterrows():
            cells = [str(v).strip() for v in row.values if pd.notna(v)]
            if cells:
                lines.append(" ".join(cells))

    return "\n".join(lines)


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_from_pdf(path)
    if ext == ".docx":
        return extract_from_docx(path)
    if ext in (".xlsx", ".xls"):
        return extract_from_excel(path)
    raise ValueError(f"Tipo de archivo no soportado: {ext}")


def parse_file(path: Path) -> Dict[str, Any]:
    raw = extract_text(path)
    lines = split_to_candidate_lines(raw)
    items = parse_lines_to_items(lines)
    return {
        "raw_text_preview": raw[:1500],
        "lines_count": len(lines),
        "items": items,
    }


