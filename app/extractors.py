from pathlib import Path
import pdfplumber
from docx import Document
import pandas as pd
import hashlib
import json

# Caché para PDFs (evitar reprocesar el mismo PDF)
_PDF_CACHE = {}

def _get_cache_key(path: Path) -> str:
    """Genera una clave de caché basada en tamaño y modificación del archivo"""
    stat = path.stat()
    key_str = f"{path.name}:{stat.st_size}:{stat.st_mtime}"
    return hashlib.md5(key_str.encode()).hexdigest()

def extract_from_pdf(path: Path, use_cache: bool = True) -> str:
    """
    Extrae texto de PDF.
    
    Args:
        path: Ruta al PDF
        use_cache: Si True, usa caché en memoria (no persiste entre ejecuciones)
    
    Returns:
        Texto extraído del PDF
    """
    if use_cache:
        cache_key = _get_cache_key(path)
        if cache_key in _PDF_CACHE:
            return _PDF_CACHE[cache_key]
    
    parts = []
    
    try:
        # Intenta extracción normal (más rápida)
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
    except Exception as e:
        # Si falla, intenta con configuración alternativa
        print(f"⚠️  PDF extraction warning: {e}")
        try:
            with pdfplumber.open(str(path), lazyload=True) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        parts.append(text)
        except Exception as e2:
            print(f"❌ PDF extraction failed: {e2}")
            return ""
    
    result = "\n".join(parts)
    
    if use_cache:
        _PDF_CACHE[cache_key] = result
    
    return result

def extract_from_docx(path: Path) -> str:
    """Extrae texto de documentos DOCX"""
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" ".join(cells))
    return "\n".join(parts)

def extract_from_excel(path: Path) -> str:
    """Extrae texto de hojas de cálculo Excel"""
    df = pd.read_excel(str(path), engine="openpyxl")
    lines = []
    for _, row in df.iterrows():
        cells = [str(v).strip() for v in row.values if pd.notna(v)]
        if cells:
            lines.append(" ".join(cells))
    return "\n".join(lines)

def extract_text(path: Path) -> str:
    """Extrae texto de PDF, DOCX o Excel"""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_from_pdf(path, use_cache=True)
    if ext == ".docx":
        return extract_from_docx(path)
    if ext in (".xlsx", ".xls"):
        return extract_from_excel(path)
    raise ValueError(f"Tipo no soportado: {ext}")

def clear_pdf_cache():
    """Limpia el caché de PDFs en memoria"""
    global _PDF_CACHE
    _PDF_CACHE.clear()
