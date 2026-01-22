from pathlib import Path
import pdfplumber
from docx import Document
import pandas as pd

def extract_from_pdf(path: Path) -> str:
    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)

def extract_from_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" ".join(cells))
    return "\n".join(parts)

def extract_from_excel(path: Path) -> str:
    df = pd.read_excel(str(path), engine="openpyxl")
    lines = []
    for _, row in df.iterrows():
        cells = [str(v).strip() for v in row.values if pd.notna(v)]
        if cells:
            lines.append(" ".join(cells))
    return "\n".join(lines)

def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_from_pdf(path)
    if ext == ".docx":
        return extract_from_docx(path)
    if ext in (".xlsx", ".xls"):
        return extract_from_excel(path)
    raise ValueError(f"Tipo no soportado: {ext}")
